# -*- coding: utf-8 -*-
"""
现病史生成模块
基于动态分段 + 完整维度的现病史模板，生成入院记录中的现病史文本。
与"首次病程·病情演变特点"的关系：现病史是完整版（现象→归纳→术语），病情演变特点是现象层简化版。
"""

from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


class PresentIllnessGenerator:
    """现病史生成器"""

    # 起病形式选项
    ONSET_FORMS = {"急": "急性起病（<2周）", "亚急": "亚急性起病（2周-3月）", "慢": "慢性起病（>3月）"}

    # 各时间段的维度字段标签
    PHASE_DIMENSIONS = [
        ("perception", "感知障碍"),
        ("thought", "思维障碍"),
        ("mood", "情感障碍"),
        ("volition", "意志行为"),
        ("cognition", "认知功能（现象层）"),
        ("social_function", "社会功能"),
    ]

    def generate(self, patient_data: Dict, output_path: str = None) -> str:
        """
        生成现病史文档

        Args:
            patient_data: 患者数据（含 present_illness 字段）
            output_path: 输出路径

        Returns:
            生成的文件路径
        """
        content = self._format_content(patient_data)

        if not output_path:
            patient_id = patient_data.get("patient_id", "unknown")
            patient_name = patient_data.get("basic_info", {}).get("name", "unknown")
            output_path = str(Path.cwd() / f"现病史_{patient_name}_{patient_id}.docx")

        self._save_to_docx(content, output_path)
        return output_path

    def _format_content(self, patient_data: Dict) -> List[Dict]:
        """
        格式化为结构化段落列表

        Returns:
            [{"title": "段落标题", "text": "段落正文", "bold": True/False}, ...]
        """
        pi = patient_data.get("history", {}).get("present_illness", {})
        if not pi:
            raise ValueError("患者数据中缺少 present_illness 字段")

        paragraphs = []

        # ===== 标题 =====
        patient_name = patient_data.get("basic_info", {}).get("name", "")
        paragraphs.append({"title": f"现病史 — {patient_name}", "text": "", "bold": True, "center": True})
        paragraphs.append({"title": "", "text": "", "bold": False})

        # ===== 信息提供者声明 =====
        informant = pi.get("informant", "")
        if informant:
            paragraphs.append({"title": "", "text": f"【信息提供者声明】", "bold": True})
            paragraphs.append({"title": "", "text": informant, "bold": False})
            paragraphs.append({"title": "", "text": "", "bold": False})

        # ===== 一、起病总述 =====
        onset = pi.get("onset", {})
        paragraphs.append({"title": "一、起病总述", "text": "", "bold": True, "heading": True})

        onset_time = onset.get("time", "")
        if onset_time:
            paragraphs.append({"title": "", "text": f"起病时间：{onset_time}", "bold": False})

        precipitant = onset.get("precipitant", "")
        if precipitant:
            paragraphs.append({"title": "", "text": f"诱因：{precipitant}", "bold": False})

        form_raw = onset.get("form", "")
        form_label = self.ONSET_FORMS.get(form_raw, form_raw)
        if form_label:
            paragraphs.append({"title": "", "text": f"起病形式：{form_label}", "bold": False})

        first_symptoms = onset.get("first_symptoms", "")
        if first_symptoms:
            paragraphs.append({"title": "", "text": f"首发症状：{first_symptoms}", "bold": False})

        paragraphs.append({"title": "", "text": "", "bold": False})

        # ===== 二、病情演变 =====
        phases = pi.get("phases", [])
        paragraphs.append({"title": "二、病情演变", "text": "", "bold": True, "heading": True})
        paragraphs.append({"title": "", "text": "", "bold": False})

        for idx, phase in enumerate(phases, 1):
            label = phase.get("label", f"时间段 {idx}")
            trigger = phase.get("trigger", "")

            paragraphs.append({"title": "", "text": f"【时间段 {idx}】{label}", "bold": True})

            if trigger:
                paragraphs.append({"title": "", "text": f"变化诱因：{trigger}", "bold": False})

            # 🆕 v4.3：各维度合并为流式段落（不再分维度列表）
            dim_texts = []
            for dim_key, dim_label in self.PHASE_DIMENSIONS:
                text = phase.get(dim_key, "")
                if text and text.strip():
                    dim_texts.append(text)

            # 本期诊疗
            treatment = phase.get("treatment", {})
            if treatment and any(treatment.values()):
                t_parts = []
                if treatment.get("date"):
                    t_parts.append(treatment["date"])
                if treatment.get("institution"):
                    t_parts.append(f"就诊于{treatment['institution']}")
                if treatment.get("diagnosis"):
                    diagnosis = treatment['diagnosis']
                    t_parts.append(f"诊断\u201c{diagnosis}\u201d")
                if treatment.get("medication"):
                    med = f"予{treatment['medication']}"
                    if treatment.get("dosage"):
                        med += f" {treatment['dosage']}"
                    t_parts.append(med)
                if treatment.get("response"):
                    t_parts.append(f"疗效：{treatment['response']}")
                if treatment.get("adverse_effects"):
                    t_parts.append(f"不良反应：{treatment['adverse_effects']}")
                if treatment.get("adherence"):
                    t_parts.append(f"依从性：{treatment['adherence']}")

                if t_parts:
                    dim_texts.append(f"本期诊疗：{'；'.join(t_parts)}")

            if dim_texts:
                flow_text = "。".join(dim_texts) + "。"
                paragraphs.append({"title": "", "text": flow_text, "bold": False})

            paragraphs.append({"title": "", "text": "", "bold": False})

        # ===== 三、发病以来一般情况 =====
        gc = pi.get("general_condition_psych", {})
        paragraphs.append({"title": "三、发病以来一般情况", "text": "", "bold": True, "heading": True})

        if gc.get("sleep"):
            paragraphs.append({"title": "", "text": f"睡眠：{gc['sleep']}", "bold": False})
        if gc.get("appetite"):
            paragraphs.append({"title": "", "text": f"饮食：{gc['appetite']}", "bold": False})
        if gc.get("bowel_bladder"):
            paragraphs.append({"title": "", "text": f"二便：{gc['bowel_bladder']}", "bold": False})
        if gc.get("weight"):
            paragraphs.append({"title": "", "text": f"体重：{gc['weight']}", "bold": False})

        # 精神科专项
        psych_specials = [
            ("impulsivity", "冲动行为"),
            ("self_harm", "自伤行为"),
            ("suicide", "自杀行为/意念"),
            ("elopement", "外走行为"),
            ("substance_use", "物质使用轨迹"),
        ]
        has_psych = any(gc.get(key, "") for key, _ in psych_specials)
        if has_psych:
            paragraphs.append({"title": "", "text": "", "bold": False})
            paragraphs.append({"title": "", "text": "精神科专项：", "bold": True})
            for key, label in psych_specials:
                val = gc.get(key, "")
                if val:
                    paragraphs.append({"title": "", "text": f"{label}：{val}", "bold": False})

        paragraphs.append({"title": "", "text": "", "bold": False})

        # ===== 四、药物汇总表 =====
        med_table = pi.get("medication_table", [])
        if med_table:
            paragraphs.append({"title": "四、药物汇总表", "text": "", "bold": True, "heading": True})
            paragraphs.append({"title": "", "text": "", "bold": False})
            paragraphs.append({
                "title": "_table_", "text": "",
                "table_data": med_table,
                "table_headers": ["时间段", "机构", "诊断", "药物", "剂量", "疗程/疗效/副作用"],
                "bold": False
            })

        return paragraphs

    def _save_to_docx(self, paragraphs: List[Dict], output_path: str):
        """保存为 Word 文档"""
        doc = Document()

        # 默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)

        for para_info in paragraphs:
            text = para_info.get("text", "")
            title = para_info.get("title", "")
            is_bold = para_info.get("bold", False)
            is_center = para_info.get("center", False)
            is_heading = para_info.get("heading", False)

            # 表格特殊处理
            if title == "_table_" and "table_data" in para_info:
                self._add_table(doc, para_info["table_headers"], para_info["table_data"])
                continue

            if not title and not text and not is_heading:
                # 空段落
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.5
                continue

            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5

            display_text = f"{title}{text}" if title else text
            run = p.add_run(display_text)
            run.font.name = '宋体'
            run.font.size = Pt(12)
            if is_bold:
                run.bold = True
            if is_center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 一级标题（一、二、三）加下划线
            if title and any(title.startswith(f"{c}、") for c in "一二三四五六七八九十"):
                run.font.size = Pt(14)
                run.underline = True

        doc.save(output_path)

    def _add_table(self, doc, headers: List[str], rows: List[Dict]):
        """添加药物汇总表"""
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10)
                    run.bold = True

        # 数据行
        for r_idx, row in enumerate(rows):
            vals = [
                row.get("period", ""),
                row.get("institution", ""),
                row.get("diagnosis", ""),
                row.get("medication", ""),
                row.get("dosage", ""),
                row.get("course_response", ""),
            ]
            for c_idx, val in enumerate(vals):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = str(val)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(10)

    def prompt_missing_fields(self, patient_data: Dict) -> str:
        """检查并提示需要填写的现病史字段"""
        prompt = "\n现病史需要以下信息：\n"
        prompt += "=" * 50 + "\n\n"

        pi = patient_data.get("history", {}).get("present_illness", {})
        if not pi:
            return prompt + "[错误] 患者数据中缺少 present_illness 字段\n"

        # 信息提供者
        if not pi.get("informant"):
            prompt += "⚠ 信息提供者声明（必填）：请注明病史来源及信息可靠性\n\n"
        else:
            prompt += f"✓ 信息提供者: {pi['informant'][:50]}...\n"

        # 起病总述
        onset = pi.get("onset", {})
        if not onset.get("time"):
            prompt += "⚠ 起病时间（必填）\n"
        if not onset.get("form"):
            prompt += "⚠ 起病形式（必填）：急性/亚急性/慢性\n"
        if not onset.get("first_symptoms"):
            prompt += "⚠ 首发症状（建议填写）\n"

        # 时间段
        phases = pi.get("phases", [])
        prompt += f"\n时间段数量: {len(phases)}\n"
        for idx, phase in enumerate(phases):
            prompt += f"\n--- 时间段 {idx+1} ({phase.get('label', '未标注')}) ---\n"
            for dim_key, dim_label in self.PHASE_DIMENSIONS:
                if not phase.get(dim_key):
                    prompt += f"  ⚠ {dim_label} 未填写\n"

        # 一般情况
        gc = pi.get("general_condition_psych", {})
        for key in ["sleep", "appetite", "impulsivity", "suicide"]:
            if not gc.get(key):
                prompt += f"⚠ 一般情况.{key} 未填写\n"

        return prompt
