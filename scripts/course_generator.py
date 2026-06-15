# -*- coding: utf-8 -*-
"""
病程记录生成模块 v5.0
关键变更：移除 format_first_course_content（语法错误+逻辑错误）
新架构：AI 直驱 → 输出结构化内容 → 代码确定性写入 Word
"""

from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re
import json
from pathlib import Path
from lxml import etree


def _set_cjk_font(run, font_name: str):
    """
    正确设置中日韩（CJK）字体
    同时设置 ASCII 和 EastAsia 字体槽位
    """
    run.font.name = font_name
    r = run._r
    rPr = r.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    if rPr is None:
        rPr = etree.SubElement(r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    if rFonts is None:
        rFonts = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', font_name)
    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', font_name)
    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font_name)
    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs', font_name)


def _set_spacing(p, before_pt: int = 0, after_pt: int = 0, line_mult: float = 1.5):
    """设置段落间距和行距"""
    pf = p.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = line_mult


class FirstCourseGenerator:
    """首次病程记录生成器 v5.0"""



    def generate_from_ai_content(self, ai_content: Dict[str, List[str]], output_path: str) -> str:
        """
        【新架构 v5.0 主入口】
        直接将 AI 生成的结构化内容写入 Word
        
        Args:
            ai_content: AI 生成的结构化内容，格式：
                {
                    "病例特点": ["（一）病例特点", "", "1. 一般特征及起病形式：...", ...],
                    "诊断依据与鉴别诊断": [...],
                    "诊疗计划": [...]
                }
            output_path: 输出 Word 文件路径
            
        Returns:
            生成的文件路径
        """
        self._save_to_docx(ai_content, output_path)
        return output_path

    def _save_to_docx(self, content: Dict[str, List[str]], output_path: str):
        """
        保存为 Word 文档（确定性写入）
        
        Args:
            content: 格式化后的内容字典
            output_path: 输出路径
        """
        doc = Document()

        # 设置默认字体（Normal 样式）
        style = doc.styles['Normal']
        style.font.name = '宋体'
        s = style._element
        rPr = s.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        if rPr is not None:
            rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
            if rFonts is None:
                rFonts = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
            rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')
        style.font.size = Pt(12)

        # 写入内容
        for section, lines in content.items():
            for line in lines:
                if not line:
                    # 空行：添加一个空段落
                    p = doc.add_paragraph()
                    _set_spacing(p, line_mult=1.5)
                    run = p.add_run("")
                    _set_cjk_font(run, '宋体')
                    run.font.size = Pt(12)
                    continue

                # 判断行类型
                is_section_title = (
                    line.startswith('（') and
                    len(line) > 1 and
                    line[1:3] in ['一）', '二）', '三）', '四）', '五）']
                )
                is_item_title = (
                    len(line) > 3 and
                    line[0].isdigit() and
                    line[1] == '.' and
                    line[2] == ' '
                )

                p = doc.add_paragraph()
                if is_section_title:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    _set_spacing(p, before_pt=6, after_pt=6, line_mult=1.5)
                    run = p.add_run(line)
                    _set_cjk_font(run, '宋体')
                    run.font.size = Pt(14)
                    run.bold = True
                elif is_item_title:
                    _set_spacing(p, line_mult=1.5)
                    run = p.add_run(line)
                    _set_cjk_font(run, '宋体')
                    run.font.size = Pt(12)
                    run.bold = True
                else:
                    _set_spacing(p, line_mult=1.5)
                    run = p.add_run(line)
                    _set_cjk_font(run, '宋体')
                    run.font.size = Pt(12)

        doc.save(output_path)
        return output_path


class DailyCourseGenerator:
    """日常病程记录生成器 v5.0 - AI直驱 + JSON中间格式
    架构：代码构建上下文 -> AI生成{markdown+snapshot} -> 代码验证保存
    """
    def _get_admission_baseline(self, patient_data: Dict) -> str:
        basic = patient_data.get("basic_info", {})
        diagnosis = patient_data.get("diagnosis", {})
        history = patient_data.get("history", {})
        exam = patient_data.get("examination", {})

        lines = []
        lines.append(f"患者：{basic.get('name','')}，{basic.get('gender','')}，{basic.get('age','')}岁")
        lines.append(f"入院日期：{basic.get('admission_date','')}")
        lines.append(f"主诉：{basic.get('chief_complaint','')}")
        lines.append(f"诊断：{diagnosis.get('primary','')}")
        pi = history.get("present_illness", {})
        if pi:
            onset = pi.get("onset", {})
            if onset:
                lines.append(f"起病形式：{onset.get('form','')}")
                lines.append(f"首发症状：{onset.get('first_symptoms','')}")
            gc = pi.get("general_condition_psych", {})
            if gc:
                parts = []
                for k, v in gc.items():
                    if v:
                        parts.append(f"{k}: {v}")
                if parts:
                    lines.append(f"入院时一般情况：{'；'.join(parts)}")
        lines.append(f"体格检查：{exam.get('physical','')[:300]}")
        lines.append(f"精神检查：{exam.get('mental','')[:500]}")
        evolution = history.get("disease_evolution", {})
        phases = evolution.get("phases", []) if isinstance(evolution, dict) else []
        if phases:
            for ph in phases:
                lines.append(f"  {ph.get('label','')}: {ph.get('phenomena','')}")
        return "\n".join(lines)

    def _build_snapshot_schema_for_ai(self) -> str:
        return (
            '"snapshot": {\n'
            '  "record_date": "YYYY-MM-DD",\n'
            '  "hospital_day": N,\n'
            '  "symptoms": {\n'
            '    "fixed": {\n'
            '      "hallucination": "string, 幻觉状态描述，若无则写无",\n'
            '      "delusion": "string, 妄想状态描述，若无则写无",\n'
            '      "mood": "string, 情绪状态描述",\n'
            '      "behavior": "string, 行为状态描述"\n'
            '    },\n'
            '    "additional": [{"name": "症状名", "status": "状态描述"}]\n'
            '  },\n'
            '  "medication": {\n'
            '    "efficacy": "string, 疗效变化描述",\n'
            '    "adverse_effects": "string, 副作用描述或无",\n'
            '    "blood_level": "string|null, 血药浓度"\n'
            '  },\n'
            '  "risks": {\n'
            '    "suicide": "低/中/高",\n'
            '    "aggression": "低/中/高",\n'
            '    "elopement": "低/中/高",\n'
            '    "somatic": "string, 躯体风险描述或无特殊"\n'
            '  }\n'
            '}'
        )

    def build_context_for_ai(
        self, patient_data: Dict, input_text: str,
        previous_record: Optional[Dict] = None
    ) -> str:
        from utils import calculate_hospital_day, format_hospital_day
        from datetime import datetime
        basic = patient_data.get("basic_info", {})
        diagnosis = patient_data.get("diagnosis", {})
        treatment = patient_data.get("treatment", {})
        admission_date = basic.get("admission_date", "")
        today_str = datetime.now().strftime("%Y-%m-%d")
        hospital_day = calculate_hospital_day(admission_date, today_str) if admission_date else 0

        patient_ctx = (
            f"患者：{basic.get('name','')}，{basic.get('gender','')}，{basic.get('age','')}岁\n"
            f"入院日期：{admission_date}\n"
            f"住院天数：第{format_hospital_day(hospital_day)}天\n"
            f"主诉：{basic.get('chief_complaint','')}\n"
            f"诊断：{diagnosis.get('primary','')}\n"
            f"治疗方案：{treatment.get('plan','')[:500]}"
        )

        baseline = self._get_admission_baseline(patient_data)

        if previous_record:
            prev_date = previous_record.get("record_date", "")
            prev_hd = previous_record.get("hospital_day", "")
            prev_snapshot = json.dumps(previous_record.get("snapshot", {}), ensure_ascii=False, indent=2)
            prev_md = previous_record.get("markdown", "")[:1000]
            previous_ctx = (
                "## 上次病程记录\n"
                f"日期：{prev_date}（住院第{prev_hd}天）\n\n"
                f"快照：\n```json\n{prev_snapshot}\n```\n\n"
                f"原文（摘要）：\n{prev_md}\n"
            )
        else:
            previous_ctx = "## 上次病程记录\n本次为首次日常病程记录，请与入院时基线状态进行对比。\n"

        comparison_guide = (
            "## 症状对比维度\n"
            "固定4维（每次必填）：\n"
            "- hallucination（幻觉）：频次、内容、类型\n"
            "- delusion（妄想）：内容、强度、信念程度\n"
            "- mood（情绪）：性质、稳定性\n"
            "- behavior（行为）：合作度、冲动性、退缩\n\n"
            "扩展槽 additional（按患者实际情况增减）：\n"
            "- 酒精依赖：戒断症状（手抖/心慌/出汗）、睡眠、渴求/复饮念头、躯体不适\n"
            "- 精神科常用：自知力变化、社交功能、药物依从性"
        )

        output_format = (
            "## 输出要求\n"
            "请输出以下 JSON 结构（不要包含任何解释性文字，只输出 JSON）：\n\n"
            "```json\n"
            "{\n"
            '  "markdown": "完整的日常病程记录Markdown文本（参考下方模板）",\n'
            + self._build_snapshot_schema_for_ai() +
            "\n}\n"
            "```\n\n"
            "## Markdown 模板\n"
            "```markdown\n"
            "# 日常病程记录\n"
            "**日期**：YYYY-MM-DD | **住院第X天** | **记录时间**：HH:MM\n\n"
            "---\n\n"
            "## 病情变化\n[当日病情变化描述，1-2段]\n\n"
            "## 精神检查\n[精神检查所见，覆盖意识/定向/接触/感知觉/思维/情感/意志/自知力]\n\n"
            "## 核心症状量化对比\n"
            "| 症状 | 入院时/上次 | 今日 | 变化 |\n"
            "|------|------------|------|------|\n"
            "| 幻觉 | ... | ... | ↑/↓/→ |\n"
            "| 妄想 | ... | ... | ↑/↓/→ |\n"
            "| 情绪 | ... | ... | ↑/↓/→ |\n"
            "| 行为 | ... | ... | ↑/↓/→ |\n"
            "[扩展症状行]\n\n"
            "## 药物反应\n"
            "- **疗效变化**：...\n"
            "- **新发副作用**：...\n"
            "- **血药浓度**：...\n\n"
            "## 风险评估\n"
            "[仅风险等级变化时出现，写变化项；无变化写各项风险等级较前无明显变化]\n\n"
            "## 处理意见\n[30-50字诊疗调整]\n"
            "```\n\n"
            "## 重要规则\n"
            "1. 若首次记录：对比列为入院时（从入院基线提取），变化列填写基线评估\n"
            "2. 固定4项的对比：必须与前次快照的数值做定量/定性对比，判断变化方向（↑恶化/↓改善/→稳定）\n"
            "3. 风险评估：仅等级变化时记录（如低→中），无变化写各项风险等级较前无明显变化\n"
            "4. 药物反应：疗效必须与具体行为指标挂钩（如幻觉频次减少而非有效）\n"
            "5. 处理意见：30-50字，合并治疗调整+病情分析+诊疗计划"
        )

        prompt = (
            "你是精神科/成瘾医学科住院医师。请根据以下信息生成今日日常病程记录。\n\n"
            f"## 患者信息\n{patient_ctx}\n\n"
            f"## 入院时基线状态\n{baseline}\n\n"
            f"{previous_ctx}\n"
            f"{comparison_guide}\n\n"
            f"## 当日临床观察\n{input_text}\n\n"
            f"{output_format}"
        )
        return prompt

    def process_ai_output(self, ai_json: Dict, patient_data: Dict) -> Dict:
        errors = []
        required_top = ["markdown", "snapshot"]
        for k in required_top:
            if k not in ai_json:
                errors.append(f"缺少顶层字段: {k}")

        if "snapshot" in ai_json:
            snap = ai_json["snapshot"]
            snap_required = ["record_date", "hospital_day", "symptoms", "medication", "risks"]
            for k in snap_required:
                if k not in snap:
                    errors.append(f"snapshot 缺少字段: {k}")
            if "symptoms" in snap:
                sym = snap["symptoms"]
                if "fixed" not in sym:
                    errors.append("symptoms 缺少 fixed 维度")
                if "additional" not in sym:
                    snap["symptoms"]["additional"] = []
            if "medication" in snap:
                med = snap["medication"]
                for k in ["efficacy", "adverse_effects", "blood_level"]:
                    if k not in med:
                        med[k] = med.get(k, None)
            if "risks" in snap:
                risks = snap["risks"]
                for k in ["suicide", "aggression", "elopement", "somatic"]:
                    if k not in risks:
                        risks[k] = risks.get(k, "未评估")

        valid = len(errors) == 0
        return {"markdown": ai_json.get("markdown", ""),
                "snapshot": ai_json.get("snapshot", {}),
                "valid": valid,
                "errors": errors}

    def _save_markdown(self, markdown: str, patient_data: Dict,
                        record_date: str) -> str:
        from pathlib import Path
        basic = patient_data.get("basic_info", {})
        name = basic.get("name", "unknown")
        pid = patient_data.get("patient_id", "unknown")
        filename = f"日常病程_{name}_{pid}_{record_date}.md"
        path = Path.cwd() / filename
        with open(str(path), "w", encoding="utf-8") as f:
            f.write(markdown)
        return str(path)